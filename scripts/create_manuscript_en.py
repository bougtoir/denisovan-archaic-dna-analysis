"""
Generate BioEssays Hypotheses manuscript (English) as .docx
Topic: Archaic introgression sharing as a tracer of ancient human migration

BioEssays format:
- Hypotheses article type (~3000-5000 words body)
- Free-form structure (not IMRaD)
- Vancouver citation style (numbered in order of appearance)
- Up to ~5-6 figures
- Abstract ~100-150 words
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from pathlib import Path

doc = Document()

# ===== Page setup =====
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Helper: add paragraph with superscript references
def add_text_with_refs(doc, text, bold=False, style_name='Normal'):
    """Parse {N} or {N-M} markers and render as superscript."""
    para = doc.add_paragraph(style=style_name)
    if bold:
        para.style.font.bold = True
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            ref_text = part[1:-1]
            run = para.add_run(ref_text)
            run.font.superscript = True
            run.font.size = Pt(10)
        else:
            run = para.add_run(part)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    return para

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_figure_with_legend(doc, fig_path, fig_num, caption):
    """Insert figure image inline followed by its legend."""
    from docx.shared import Inches
    fig_file = Path(fig_path)
    if fig_file.exists():
        para_img = doc.add_paragraph()
        para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = para_img.add_run()
        run_img.add_picture(str(fig_file), width=Inches(6.0))
    else:
        doc.add_paragraph(f'[Figure {fig_num}: {fig_path} not found]')
    para = doc.add_paragraph()
    run = para.add_run(f'Figure {fig_num}. ')
    run.bold = True
    run.font.size = Pt(11)
    run = para.add_run(caption)
    run.font.size = Pt(11)
    return para

# ===== TITLE PAGE =====
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run(
    'Archaic Introgression Sharing as a Tracer of Ancient Human Migration:\n'
    'A Bivariate Approach Using Neanderthal and Denisovan DNA Signatures'
)
run.font.size = Pt(16)
run.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()

# Authors
author_para = doc.add_paragraph()
author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author_para.add_run('Tatsuki Onishi')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run2 = author_para.add_run('1')
run2.font.superscript = True
run2.font.size = Pt(10)
run3 = author_para.add_run('*')
run3.font.superscript = True
run3.font.size = Pt(10)

doc.add_paragraph()

# Affiliation
affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = affil.add_run('1 [Affiliation to be added]')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

doc.add_paragraph()

# Corresponding author
corr = doc.add_paragraph()
corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = corr.add_run('*Corresponding author: Tatsuki Onishi (bougtoir@gmail.com)')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

doc.add_paragraph()

# Article type
atype = doc.add_paragraph()
atype.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = atype.add_run('Article type: Hypotheses')
run.font.size = Pt(11)
run.bold = True

# Running title
rtitle = doc.add_paragraph()
rtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = rtitle.add_run('Running title: Archaic DNA as migration tracer')
run.font.size = Pt(10)
run.italic = True

doc.add_paragraph()

# Keywords
kw = doc.add_paragraph()
run = kw.add_run('Keywords: ')
run.bold = True
run2 = kw.add_run(
    'archaic introgression, Denisovan, Neanderthal, human migration, '
    'population genetics, introgression sharing, Wallace Line, admixture'
)

doc.add_page_break()

# ===== ABSTRACT =====
add_heading(doc, 'Abstract', level=1)

abstract_text = (
    'Archaic human DNA — inherited from Neanderthals and Denisovans — persists in '
    'modern human genomes in population-specific patterns shaped by ancient migration '
    'routes. We hypothesize that pairwise correlations of archaic introgression '
    'frequency profiles between populations can serve as an independent tracer of '
    'past long-distance human movement, complementing conventional approaches based on '
    'genome-wide ancestry proportions. '
    'Using publicly available hmmix introgression segments from 3,134 individuals '
    'across 66 populations (1000 Genomes + HGDP), we demonstrate that a multiple '
    'regression model incorporating geographic distance, recent admixture proportions, '
    'and continental grouping explains 51% of the variance in Neanderthal sharing '
    'correlations (R{2} = 0.510) and 50% for Denisovan (R{2} = 0.495). Crucially, '
    'Denisovan introgression creates a sharp biogeographic boundary coinciding with '
    'the Wallace Line, while Neanderthal sharing residuals highlight trans-Pacific '
    'connections consistent with Beringian migration. We propose specific testable '
    'predictions for validation with forthcoming ancient genome data.'
)
add_text_with_refs(doc, abstract_text)

doc.add_page_break()

# ===== INTRODUCTION =====
add_heading(doc, 'Introduction', level=1)

intro_paras = [
    ('The genomes of non-African modern humans carry detectable traces of archaic '
     'human DNA acquired through interbreeding events during the Out-of-Africa '
     'dispersal. Neanderthal ancestry accounts for approximately 1-2% of the genomes '
     'of all non-African populations,{1,2} while Denisovan ancestry shows a striking '
     'geographic gradient: 3-5% in Oceanian populations (Papuans, Australian '
     'Aboriginals, Philippine Ayta), yet only ~0.06% in East Asians and ~0.02% in '
     'Europeans.{3,4} This asymmetry reflects the geography of archaic admixture: '
     'Neanderthal interbreeding occurred once in western Asia (affecting all non-African '
     'lineages), whereas major Denisovan admixture occurred in Southeast Asia, '
     'primarily benefiting ancestors of Oceanian populations.{5}'),

    ('The spatial distribution of archaic DNA has been used to infer properties of '
     'the admixture events themselves — their timing, number, and geographic '
     'location.{4,6} However, the downstream redistribution of archaic segments '
     'through subsequent human migration has received less systematic attention. '
     'Quilodran et al. demonstrated that spatial gradients of Neanderthal ancestry '
     'across Eurasia could be decomposed into three historically meaningful migration '
     'waves: the initial Out-of-Africa dispersal, the Neolithic farming expansion, '
     'and the Steppe pastoralist migration.{7} Their approach, however, focused '
     'exclusively on Neanderthal DNA and modeled aggregate ancestry proportions rather '
     'than the fine-grained sharing of specific introgressed segments.'),

    ('Here, we propose and provide initial evidence for a complementary approach: '
     'using pairwise correlations of introgression frequency profiles — the degree to '
     'which two populations share the same archaic DNA segments at the same genomic '
     'positions — as an independent tracer of shared migration history. Our approach '
     'extends beyond Quilodran et al. in three key ways: (1) we jointly analyze both '
     'Neanderthal and Denisovan introgression, exploiting their complementary geographic '
     'signatures; (2) we operate at the level of segment sharing rather than aggregate '
     'proportions, capturing population-specific introgression patterns; and (3) we '
     'explicitly model confounding factors (recent admixture and continental grouping) '
     'to isolate residual sharing signals indicative of unexpected historical connections.'),
]

for para_text in intro_paras:
    add_text_with_refs(doc, para_text)

# ===== THE HYPOTHESIS =====
add_heading(doc, 'The Hypothesis', level=1)

hyp_paras = [
    ('We hypothesize that the pairwise correlation of archaic introgression frequency '
     'profiles between human populations — after correction for geographic distance, '
     'recent admixture, and shared continental ancestry — contains residual signal '
     'that reflects ancient migration events not captured by conventional population '
     'genetic summary statistics. Specifically:'),

    ('(1) Neanderthal introgression sharing, being shared by all non-African lineages, '
     'acts as a "common-mode" tracer whose residual variation reflects post-admixture '
     'population movements (analogous to a shared baseline modulated by drift and migration).'),

    ('(2) Denisovan introgression, with its extreme geographic concentration in Oceania, '
     'provides a "differential-mode" tracer that marks the Wallace Line crossing and can '
     'detect any subsequent gene flow between Oceanian and non-Oceanian populations.'),

    ('(3) The combination of these two signals — one broadly distributed, one sharply '
     'localized — creates a bivariate "fingerprint" that can distinguish migration '
     'routes more effectively than either signal alone.'),
]

for para_text in hyp_paras:
    add_text_with_refs(doc, para_text)

# ===== SUPPORTING EVIDENCE =====
add_heading(doc, 'Supporting Evidence from Publicly Available Data', level=1)

add_heading(doc, 'Data and analytical framework', level=2)

evidence_paras = [
    ('To evaluate this hypothesis, we reanalyzed archaic introgression segments '
     'detected by hmmix{8} in 3,134 individuals from 66 populations (1000 Genomes '
     'Project + Human Genome Diversity Project), publicly available via '
     'Zenodo (record 14136628). We binned the genome into 500 kb windows, computed '
     'the frequency of archaic (Neanderthal or Denisovan) introgression in each bin '
     'for each population, and calculated Pearson correlations between all population '
     'pairs. This yielded a 66 x 66 sharing matrix for each archaic source.'),

    ('To address known confounders, we fitted a multiple regression model:\n\n'
     '    Sharing_ij = beta_0 + beta_1 * Distance_ij + beta_2 * MaxAdmixEur_ij + '
     'beta_3 * SameContinent_ij + epsilon_ij\n\n'
     'where MaxAdmixEur captures the maximum European admixture fraction in either '
     'population of the pair (addressing the inflation of sharing correlations in '
     'recently admixed American populations such as PUR, CLM, MXL, and PEL), and '
     'SameContinent is a binary indicator for shared continental grouping.'),
]

for para_text in evidence_paras:
    add_text_with_refs(doc, para_text)

add_heading(doc, 'Confounding correction substantially improves model fit', level=2)

correction_paras = [
    ('The uncorrected model (geographic distance only) explained 23.5% of the variance '
     'in Neanderthal sharing correlations and 21.7% for Denisovan. After incorporating '
     'admixture and continental grouping, the corrected model explained 51.0% '
     '(Neanderthal, 95% CI: 48.3-53.8%) and 49.5% (Denisovan, 95% CI: 46.9-52.2%), '
     'as confirmed by 5,000 bootstrap resamples (Figure 1). All three covariates '
     'contributed significantly: geographic distance (beta = -0.016 per 1000 km, '
     'p < 10{-15}), European admixture fraction (beta = 0.42, p < 10{-15}), and '
     'same-continent indicator (beta = 0.30, p < 10{-15}). The Mantel test confirmed '
     'that the distance-sharing relationship remains highly significant even when '
     'restricted to non-admixed population pairs (Neanderthal: r = -0.62, p = 0.0001; '
     'Denisovan: r = -0.58, p = 0.0001).'),

    ('The partial correlation between sharing and distance, after controlling for '
     'admixture and continental grouping, decreased from r = -0.485 to r = -0.311 '
     'for Neanderthal and from r = -0.466 to r = -0.300 for Denisovan. This indicates '
     'that approximately 36% of the raw distance-sharing correlation was attributable to '
     'confounding by continental structure and recent admixture — a substantial '
     'correction that validates the necessity of this adjustment.'),
]

for para_text in correction_paras:
    add_text_with_refs(doc, para_text)

# Figure 1 (inline)
add_figure_with_legend(doc, 'figures/fig1_sharing_vs_distance.png', 1,
    'Archaic DNA sharing correlation vs. geographic distance. '
    '(A) Neanderthal segment sharing: blue dots indicate same-continent pairs, '
    'red dots indicate cross-continent pairs, and orange triangles indicate pairs '
    'involving recently admixed populations (PUR, CLM, MXL, PEL). Dashed lines show '
    'the corrected regression model for cross-continental (black) and same-continental '
    '(blue) pairs. Key outlier pairs are annotated. '
    '(B) Denisovan segment sharing: purple diamonds indicate Oceania-involved pairs, '
    'which cluster as a distinct group with negative or near-zero correlations with '
    'non-Oceanian populations, reflecting the Denisovan admixture boundary at the '
    'Wallace Line. '
    'Data: hmmix introgression segments (Zenodo:14136628), 66 populations, '
    '3,134 individuals, 500 kb bins, corrected for admixture and continental grouping.')

add_heading(doc, 'Denisovan DNA delineates the Wallace Line boundary', level=2)

# Wallace para 1 — cites Figure 2
add_text_with_refs(doc,
    'The most striking feature of the Denisovan sharing data is the complete '
    'separation of Oceanian populations (PapuanHighlands, PapuanSepik, Bougainville) '
    'from all other populations (Figure 1B, Figure 2). Oceanian populations share '
    'high Denisovan introgression correlations with each other (r = 0.75-0.80) but '
    'show near-zero or negative correlations with non-Oceanian populations (r = -0.15 '
    'to 0.10). This binary pattern is far sharper than the gradual distance-decay '
    'observed for Neanderthal DNA and corresponds precisely to the Wallace Line — '
    'the biogeographic boundary between the Asian and Australian continental shelves.')

# Figure 2 (inline — immediately after first citation)
add_figure_with_legend(doc, 'figures/fig2_sharing_heatmap.png', 2,
    'Pairwise archaic DNA segment sharing heatmap for 31 key populations. '
    'Left: Neanderthal sharing shows a continuous gradient from high (within-continent) '
    'to low (cross-continent) correlations, with notable cross-continental '
    'connectivity in admixed American populations. Right: Denisovan sharing reveals '
    'a binary structure — Oceanian populations (bottom-right block) form an isolated '
    'cluster with high intra-group sharing but near-zero sharing with all other populations. '
    'Population labels are colored by continental region.')

# Wallace para 2 — cites Figure 3
add_text_with_refs(doc,
    'This finding has two implications. First, the Denisovan admixture event(s) in '
    'Oceania involved a distinct Denisovan population or lineage whose introgressed '
    'segments do not overlap with the low-level Denisovan ancestry (~0.06%) found in '
    'East Asians.{4} This is consistent with Jacobs et al., who identified at least '
    'three distinct Denisovan ancestries in Oceanian populations.{5} Second, the '
    'sharpness of this boundary indicates that post-admixture gene flow across the '
    'Wallace Line has been extremely limited — the archaic DNA "stamp" has been '
    'preserved through ~45,000 years of population history (Figure 3).')

# Figure 3 (inline — immediately after first citation)
add_figure_with_legend(doc, 'figures/fig3_minard_migration.png', 3,
    'Minard-style flow visualization of Homo sapiens Out-of-Africa migration '
    'with archaic admixture events. Band width represents relative population size '
    '(reflecting decreasing genetic diversity with distance from Africa). Stars mark '
    'admixture events: Neanderthal admixture ~47,000 years ago (affecting all '
    'non-African lineages), Denisovan admixture 1 ~45,000 years ago (Oceanian '
    'lineage, 3-5%), and Denisovan admixture 2 ~30,000 years ago (East Asian '
    'lineage, ~0.06%). Yellow and pink shaded regions indicate Neanderthal and '
    'Denisovan geographic ranges, respectively.')

add_heading(doc, 'Neanderthal residuals highlight trans-Pacific connections', level=2)

trans_pacific = [
    ('Among non-admixed population pairs, the largest positive residuals from the '
     'corrected Neanderthal model involve Middle East-Europe pairs '
     '(Palestinian-TSI: residual +0.395; IBS-Palestinian: +0.386), reflecting the '
     'close genetic relationship between these regions despite their assignment to '
     'different continental categories. More informative are the cross-continental '
     'outliers: among admixed pairs, East Asian-Peruvian connections (KHV-PEL: '
     'residual +0.353; CHS-PEL: +0.346; CHB-PEL: +0.334) show elevated sharing '
     'over 16,000-19,000 km, consistent with the Beringian migration route.'),

    ('While individual outlier pairs do not reach statistical significance after '
     'FDR correction (q > 0.10 for all pairs), the consistent directionality of '
     'the top residuals — preferentially connecting East Asian and American populations '
     'rather than, for example, European and Oceanian populations — supports the '
     'hypothesis that Neanderthal segment sharing retains a signal of the Beringian '
     'crossing. We note that the absence of individually significant outliers is '
     'expected given the modest sample size per population (median = 20 individuals) '
     'and the conservative nature of genome-wide permutation tests with 2,145 pair '
     'comparisons.'),
]

for para_text in trans_pacific:
    add_text_with_refs(doc, para_text)

# ===== DIFFERENTIATION FROM PRIOR WORK =====
add_heading(doc, 'Relationship to Prior Approaches', level=1)

diff_paras = [
    ('Our hypothesis builds on, but differs fundamentally from, several prior '
     'approaches to using archaic DNA for demographic inference:'),

    ('Quilodran et al.{7} analyzed spatial gradients of aggregate Neanderthal ancestry '
     'proportions across Eurasian populations and recovered three migration waves. '
     'Their approach operates on mean ancestry fractions, treating Neanderthal DNA '
     'as a scalar value per population. Our approach instead compares the full '
     'introgression frequency profile — a vector of ~6,000 genomic bins — between '
     'population pairs, capturing which specific segments are shared rather than '
     'just how much total archaic DNA is present. This distinction matters because '
     'two populations can have identical mean Neanderthal ancestry (~1.4%) yet share '
     'different subsets of introgressed segments, reflecting different post-admixture '
     'drift histories.'),

    ('Petr et al.{9} analyzed Neanderthal ancestry patterns in ancient European '
     'genomes, evaluating whether introgression levels have declined over time. '
     'Their approach leverages ancient DNA to provide temporal resolution but requires '
     'well-preserved samples, which are scarce for tropical and southern-hemisphere '
     'populations. Our spatial approach using modern genomes can access populations '
     '(Oceanian, Southeast Asian, South American) for which ancient DNA is largely '
     'unavailable.'),

    ('Mao et al.{10} traced Denisovan segments in European populations to indirect '
     'gene flow via South Asian intermediaries. Our analysis recapitulates and extends '
     'this finding: the Denisovan sharing heatmap (Figure 2, right panel) shows '
     'non-trivial sharing between Central/South Asian and European populations, '
     'consistent with indirect Denisovan gene flow through the Eurasian interior.'),
]

for para_text in diff_paras:
    add_text_with_refs(doc, para_text)

# ===== TESTABLE PREDICTIONS =====
add_heading(doc, 'Testable Predictions', level=1)

pred_paras = [
    ('Our hypothesis generates several specific, falsifiable predictions:'),

    ('Prediction 1: Ancient genomes from Beringia. If high-coverage genomes become '
     'available from the Beringian Standstill period (~24,000-18,000 years ago), '
     'their Neanderthal introgression profiles should show elevated sharing '
     'correlations with both East Asian (CHB, CHS, JPT) and Native American '
     '(Karitiana, Surui, Pima) modern populations, bridging the gap that currently '
     'separates these groups in our sharing matrix. The predicted sharing correlation '
     'should be r > 0.6 with both groups, higher than the r ~ 0.3-0.4 currently '
     'observed between modern East Asian and South American populations.'),

    ('Prediction 2: Denisovan segment phylogeny. If the distinct Denisovan lineages '
     'identified by Jacobs et al.{5} can be classified at the segment level, '
     'the Oceanian-type segments should be completely absent in Japanese and Korean '
     'populations (consistent with the JEWEL study finding of only 1.47 Mb Denisovan '
     'DNA in Japanese genomes{11}), while the East Asian-type segments should show a '
     'north-south gradient within East Asia consistent with a Southeast Asian admixture '
     'source.'),

    ('Prediction 3: Population Y signal. Skoglund et al. identified ~2% '
     'Australasian-like ancestry in Amazonian populations (Surui, Karitiana).{12} '
     'If this ancestry derives from a population with Oceanian-type Denisovan DNA, '
     'then the Denisovan sharing correlation between Surui/Karitiana and '
     'Papuan populations should be detectably higher than between other Native '
     'American populations and Papuans. Current sample sizes (n ~ 20 per population) '
     'may be insufficient, but targeted sequencing of ~100 individuals per group '
     'should provide adequate power.'),

    ('Prediction 4: Post-Columbian admixture as a positive control. The high '
     'Neanderthal sharing observed between admixed American populations (PUR, CLM) '
     'and European populations should disappear entirely when conditioning on the '
     'European ancestry component of the admixed individuals. This provides a '
     'methodological validation: if our approach correctly attributes this elevated '
     'sharing to recent admixture rather than ancient connection, it confirms that '
     'the corrected residuals reflect genuinely ancient signals.'),
]

for para_text in pred_paras:
    add_text_with_refs(doc, para_text)

# Figure 4 (inline — after Testable Predictions, before Limitations)
add_text_with_refs(doc,
    'The global distribution of archaic DNA — with Neanderthal ancestry broadly '
    'distributed across non-African populations and Denisovan ancestry sharply '
    'concentrated in Oceania — is summarized in Figure 4. This bivariate '
    'representation captures the complementary geographic signatures that '
    'underpin our hypothesis: the "common-mode" Neanderthal signal '
    '(circle size) vs. the "differential-mode" Denisovan signal (color intensity).')

add_figure_with_legend(doc, 'figures/fig4_bivariate_world_map.png', 4,
    'Global distribution of archaic human DNA: bivariate world map. Circle size '
    'represents Neanderthal DNA proportion (0.08-1.8%). Color intensity represents '
    'Denisovan DNA proportion (0.02-3.5%). The sharp transition from low '
    '(yellow/pale) to high (red/dark) Denisovan proportions between mainland '
    'Southeast Asia and island Melanesia visualizes the Wallace Line boundary. '
    'Japanese populations (small yellow circles) have substantial Neanderthal '
    'ancestry (~1.4%) but minimal Denisovan ancestry (~0.06%), placing them in '
    'the "high-Neanderthal, low-Denisovan" quadrant shared with other continental '
    'East Asian populations.')

# ===== LIMITATIONS =====
add_heading(doc, 'Limitations and Caveats', level=1)

limit_paras = [
    ('Several limitations of the current analysis should be acknowledged. First, '
     'the Pearson correlation of binned introgression frequencies is a '
     'computationally simple but statistically limited metric. It does not account '
     'for linkage disequilibrium between adjacent bins, nor does it distinguish '
     'between identity-by-descent and identity-by-state of archaic segments. More '
     'sophisticated approaches — such as phylogenetic analysis of shared segment '
     'haplotypes or hidden Markov models of segment co-inheritance — would provide '
     'stronger evidence.'),

    ('Second, our confounding correction uses a binary continental grouping indicator '
     'rather than a continuous measure of genetic distance (e.g., F_ST). Published '
     'F_ST values for all 66 population pairs in our dataset are not uniformly '
     'available, and computing them from the hmmix segment data would be circular. '
     'Future work should incorporate independent estimates of genome-wide genetic '
     'distance as covariates.'),

    ('Third, the sample sizes per population (median ~20 individuals) limit '
     'statistical power for detecting subtle sharing signals, particularly for '
     'Denisovan DNA, which constitutes a small fraction of total introgression. '
     'The absence of individually significant outlier pairs after FDR correction '
     'does not refute the hypothesis but indicates that larger sample sizes are '
     'needed for definitive tests.'),
]

for para_text in limit_paras:
    add_text_with_refs(doc, para_text)

# ===== CONCLUSION =====
add_heading(doc, 'Conclusions and Outlook', level=1)

conclusion_paras = [
    ('We have proposed that pairwise archaic introgression sharing — the correlation '
     'of introgression frequency profiles between populations — can serve as an '
     'independent tracer of ancient human migration. Using publicly available data, '
     'we demonstrated that this metric captures biologically meaningful signals: '
     'the Denisovan Wallace Line boundary, the Neanderthal-mediated Beringian '
     'connection, and the confounding effect of post-Columbian admixture. After '
     'correcting for these confounders, the signal persists but at reduced statistical '
     'power, indicating that larger datasets and more sophisticated analytical methods '
     'are needed to fully exploit this approach.'),

    ('The key advantage of our bivariate framework — jointly using Neanderthal and '
     'Denisovan signatures — is its ability to distinguish migration routes that are '
     'degenerate when viewed through a single archaic lens. A population with 1.4% '
     'Neanderthal and 0.06% Denisovan DNA (Japanese) has a fundamentally different '
     'migration history from one with 1.8% Neanderthal and 3.5% Denisovan (Papuan), '
     'even though both carry substantial archaic ancestry. As ancient genome sequencing '
     'extends to underrepresented regions — particularly Beringia, island Southeast '
     'Asia, and the Pacific — the testable predictions outlined here will become '
     'addressable, potentially resolving longstanding questions about the routes and '
     'timing of human dispersal across the planet.'),
]

for para_text in conclusion_paras:
    add_text_with_refs(doc, conclusion_paras[0] if para_text == conclusion_paras[0] else para_text)

# ===== DATA AVAILABILITY =====
add_heading(doc, 'Data Availability', level=1)
da = doc.add_paragraph(
    'All analysis scripts, derived data, and figures are available at '
    'https://github.com/bougtoir/denisovan-archaic-dna-analysis. '
    'Source data: hmmix archaic introgression segments, Zenodo record 14136628.'
)

# ===== ACKNOWLEDGMENTS =====
add_heading(doc, 'Acknowledgments', level=1)
doc.add_paragraph('[To be added]')

# ===== CONFLICT OF INTEREST =====
add_heading(doc, 'Conflict of Interest', level=1)
doc.add_paragraph('The author declares no conflict of interest.')

# ===== REFERENCES =====
add_heading(doc, 'References', level=1)

references = [
    '1. Sankararaman S, Mallick S, Dannemann M, et al. The genomic landscape of '
    'Neanderthal ancestry in present-day humans. Nature. 2014;507(7492):354-357.',

    '2. Prufer K, de Filippo C, Grote S, et al. A high-coverage Neandertal genome '
    'from Vindija Cave in Croatia. Science. 2017;358(6363):655-658.',

    '3. Sankararaman S, Mallick S, Patterson N, Reich D. The combined landscape of '
    'Denisovan and Neanderthal ancestry in present-day humans. Curr Biol. '
    '2016;26(9):1241-1247.',

    '4. Reich D, Green RE, Kircher M, et al. Genetic history of an archaic hominin '
    'group from Denisova Cave in Siberia. Nature. 2010;468(7327):1053-1060.',

    '5. Jacobs GS, Hudjashov G, Saag L, et al. Multiple deeply divergent Denisovan '
    'ancestries in Papuans. Cell. 2019;177(4):1010-1021.',

    '6. Vernot B, Akey JM. Resurrecting surviving Neandertal lineages from modern '
    'human genomes. Science. 2014;343(6174):1017-1021.',

    '7. Quilodr\u00e1n CS, Rio J, Tsoupas A, Currat M. Past human expansions shaped '
    'the spatial pattern of Neanderthal ancestry. Sci Adv. 2023;9(42):eadg9817.',

    '8. Skov L, Hui R, Shchur V, et al. Detecting archaic introgression using '
    'an unadmixed outgroup. PLoS Genet. 2018;14(9):e1007641.',

    '9. Petr M, P\u00e4\u00e4bo S, Kelso J, Vernot B. Limits of long-term selection '
    'against Neandertal introgression. Proc Natl Acad Sci U S A. 2019;116(5):1639-1644.',

    '10. Mao X, Zhang H, Qiao S, et al. The deep population history of northern '
    'East Asia from the Late Pleistocene to the Holocene. Cell. 2021;184(12):3256-3266.',

    '11. Liu X, Koyama S, Tomizuka K, et al. Decoding triancestral origins, '
    'archaic introgression, and natural selection in the Japanese population by '
    'whole-genome sequencing. Sci Adv. 2024;10(16):eadi8419.',

    '12. Skoglund P, Mallick S, Bortolini MC, et al. Genetic evidence for two '
    'founding populations of the Americas. Nature. 2015;525(7567):104-108.',
]

for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)

# ===== Save =====
outpath = Path('docs/manuscript_bioessays_en.docx')
doc.save(str(outpath))
print(f"Manuscript saved to {outpath}")
print(f"Word count (approximate): ~4000 words")
