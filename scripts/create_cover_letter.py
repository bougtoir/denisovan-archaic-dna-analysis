"""
Generate BioEssays cover letter as .docx
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import datetime

doc = Document()

# Page setup
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Date
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = date_para.add_run(datetime.date.today().strftime('%B %d, %Y'))

doc.add_paragraph()

# Addressee
addr = doc.add_paragraph()
run = addr.add_run('The Editor')
doc.add_paragraph('BioEssays')
doc.add_paragraph('Wiley')

doc.add_paragraph()

# Subject
subj = doc.add_paragraph()
run = subj.add_run('Re: Submission of Hypotheses article')
run.bold = True

doc.add_paragraph()

# Salutation
doc.add_paragraph('Dear Editor,')

doc.add_paragraph()

# Body paragraphs
paras = [
    ('We are pleased to submit the enclosed manuscript entitled '
     '"Archaic Introgression Sharing as a Tracer of Ancient Human Migration: '
     'A Bivariate Approach Using Neanderthal and Denisovan DNA Signatures" '
     'for consideration as a Hypotheses article in BioEssays.'),

    ('This manuscript proposes that pairwise correlations of archaic '
     'introgression frequency profiles between human populations can serve '
     'as an independent tracer of past long-distance migration. Using publicly '
     'available hmmix introgression segments from 3,134 individuals across '
     '66 populations (1000 Genomes + HGDP), we demonstrate that a multiple '
     'regression model incorporating geographic distance, recent admixture '
     'proportions, and continental grouping explains over 50% of the variance '
     'in both Neanderthal and Denisovan sharing correlations. We identify '
     'two key geographic signatures: a sharp Denisovan boundary at the Wallace '
     'Line and Neanderthal sharing residuals consistent with Beringian migration.'),

    ('Our work extends beyond the recent study by Quilodr\u00e1n et al. (Sci Adv, '
     '2023) in three important ways: (1) we jointly analyze both Neanderthal '
     'and Denisovan introgression as complementary tracers; (2) we operate at '
     'the level of segment sharing rather than aggregate ancestry proportions; '
     'and (3) we explicitly model and correct for confounding factors including '
     'recent admixture in American populations and shared continental ancestry. '
     'We also provide four specific, testable predictions for validation with '
     'forthcoming ancient genome data.'),

    ('We believe this hypothesis-driven manuscript is well suited to the '
     'Hypotheses format of BioEssays, as it proposes a novel conceptual '
     'framework supported by reanalysis of existing public data, without '
     'requiring original experimental work. The bivariate approach we outline '
     'opens new avenues for using archaic DNA as an independent line of '
     'evidence in reconstructing human migration history.'),

    ('All analysis code and derived data are publicly available on GitHub. '
     'The manuscript has not been published previously and is not under '
     'consideration elsewhere. All authors have approved the submission.'),
]

for text in paras:
    doc.add_paragraph(text)

doc.add_paragraph()

# Closing
doc.add_paragraph('We look forward to your consideration.')

doc.add_paragraph()
doc.add_paragraph('Sincerely,')

doc.add_paragraph()

# Author
author = doc.add_paragraph()
run = author.add_run('Tatsuki Onishi')
run.bold = True

affil = doc.add_paragraph('[Affiliation to be added]')
for run in affil.runs:
    run.font.size = Pt(10)

email = doc.add_paragraph('Email: bougtoir@gmail.com')
for run in email.runs:
    run.font.size = Pt(10)

# Save
outpath = Path('docs/cover_letter_bioessays.docx')
doc.save(str(outpath))
print(f"Cover letter saved to {outpath}")
